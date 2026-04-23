from docx import Document
import sys

path = 'D:\\AI编程\\claude测试\\VS code claude\\data\\other books\\芒格之道.docx'
doc = Document(path)

count = 0
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t:
        continue
    count += 1
    if count <= 80:
        print(f'[{i}] {p.style.name!r:25s} | {t[:80]}')

print(f'\nTotal non-empty paragraphs: {count}')
print(f'Total paragraphs: {len(doc.paragraphs)}')
